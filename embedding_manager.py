from chroma_config import CHROMA_SETTINGS
from langchain_community.embeddings import OllamaEmbeddings
from langchain.vectorstores import Chroma
from langchain.schema import Document
from typing import List, Optional, Tuple
import os
import docx
import pypdf
from pathlib import Path

class EmbeddingManager:
    def __init__(self):
        os.makedirs(CHROMA_SETTINGS["persist_directory"], exist_ok=True)

        self.embedding_model = OllamaEmbeddings(
            model="all-minilm:latest"
        )

        self.db = Chroma(
            persist_directory=CHROMA_SETTINGS["persist_directory"],
            embedding_function=self.embedding_model,
            collection_name=CHROMA_SETTINGS["collection_name"]
        )

    def add_documents(self, documents: List[Document]):
        if not isinstance(documents, list) or not all(isinstance(doc, Document) for doc in documents):
            raise ValueError("Input must be a list of Document objects")
        self.db.add_documents(documents)
        self.db.persist()

    def similarity_search(self, query: str, k: int = 3) -> List[Document]:
        return self.db.similarity_search(query, k=k)

    def similarity_search_with_score(self, query: str, k: int = 3) -> List[Tuple[Document, float]]:
        return self.db.similarity_search_with_score(query, k=k)

    def add_text(self, text: str, source: Optional[str] = None):
        doc = Document(page_content=text, metadata={"source": source})
        self.add_documents([doc])

    def load_and_add_file(self, file_path: str):
        extension = Path(file_path).suffix.lower()
        if extension == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                text = file.read()
        elif extension == ".pdf":
            with open(file_path, "rb") as file:
                reader = pypdf.PdfReader(file)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif extension == ".docx":
            doc = docx.Document(file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

        self.add_text(text, source=file_path)

    def debug_status(self):
        return {
            "persist_directory": CHROMA_SETTINGS["persist_directory"],
            "collection": CHROMA_SETTINGS["collection_name"],
            "doc_count": self.db._collection.count(),
            "model": "all-minilm:latest"
        }
